"""
File parsing utilities for resume analyzer
Supports PDF, DOC, DOCX, and TXT files
"""

import io
import pdfplumber
from docx import Document
import streamlit as st


def parse_uploaded_file(uploaded_file):
    """
    Parse uploaded file and extract text content
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content or None if parsing failed
    """
    try:
        file_type = uploaded_file.type
        file_name = uploaded_file.name
        
        if file_type == "text/plain":
            # Handle TXT files
            return str(uploaded_file.read(), "utf-8")
            
        elif file_type == "application/pdf":
            # Handle PDF files
            return parse_pdf(uploaded_file)
            
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"]:
            # Handle DOCX and DOC files
            return parse_docx(uploaded_file)
            
        else:
            st.error(f"Unsupported file type: {file_type}")
            return None
            
    except Exception as e:
        st.error(f"Error parsing file {uploaded_file.name}: {str(e)}")
        return None


def parse_pdf(uploaded_file):
    """
    Extract text from PDF file using pdfplumber
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    try:
        # Create a BytesIO object from the uploaded file
        pdf_bytes = io.BytesIO(uploaded_file.read())
        
        text_content = []
        
        with pdfplumber.open(pdf_bytes) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        return "\n".join(text_content)
        
    except Exception as e:
        st.error(f"Error parsing PDF: {str(e)}")
        return None


def parse_docx(uploaded_file):
    """
    Extract text from DOCX file using python-docx
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    try:
        # Create a BytesIO object from the uploaded file
        docx_bytes = io.BytesIO(uploaded_file.read())
        
        # Load the document
        doc = Document(docx_bytes)
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        return "\n".join(text_content)
        
    except Exception as e:
        st.error(f"Error parsing DOCX: {str(e)}")
        return None


def validate_resume_content(text_content):
    """
    Validate that the extracted text looks like a resume
    
    Args:
        text_content (str): Extracted text content
        
    Returns:
        dict: Validation results with warnings and suggestions
    """
    if not text_content or len(text_content.strip()) < 50:
        return {
            "valid": False,
            "warning": "The extracted text is too short. Please ensure your resume contains sufficient content.",
            "suggestions": [
                "Check if the file is corrupted",
                "Try uploading a different format (PDF, DOCX, TXT)",
                "Use the 'Paste Resume Text' option instead"
            ]
        }
    
    # Check for common resume indicators
    resume_indicators = [
        "experience", "education", "skills", "work", "job", "position",
        "university", "college", "degree", "certification", "project",
        "email", "phone", "contact", "summary", "objective"
    ]
    
    text_lower = text_content.lower()
    found_indicators = [indicator for indicator in resume_indicators if indicator in text_lower]
    
    if len(found_indicators) < 3:
        return {
            "valid": True,
            "warning": "This doesn't look like a typical resume format.",
            "suggestions": [
                "Ensure the document contains sections like Experience, Education, Skills",
                "Include contact information and work history",
                "Add relevant keywords for better analysis"
            ]
        }
    
    return {
        "valid": True,
        "warning": None,
        "suggestions": []
    }


def get_file_info(uploaded_file):
    """
    Get information about the uploaded file
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        dict: File information
    """
    return {
        "name": uploaded_file.name,
        "size": uploaded_file.size,
        "type": uploaded_file.type,
        "size_mb": round(uploaded_file.size / (1024 * 1024), 2)
    }
